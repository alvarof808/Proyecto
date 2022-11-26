import docx
import random
from random import sample
import re
from docx import Document

doc = docx.Document('1.docx')

# Funcion para modificar el archivos .docx
def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

# Funcion para obtener todo el texto del documento .docx
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

# Funcion para obtener los parrafos independientes
def getParrafo(filename):
    doc = docx.Document(filename)
    parrafo = []
    x = len(doc.paragraphs)
    for y in range(0, x):
        parrafo.append(doc.paragraphs[y].text)
    return parrafo  
  
w  = getParrafo('1.docx') 

# Funcion para separar las palabras del parrafo
def getPalabras(listado):
    palabras = listado.split(sep=' ')
    return palabras
p = w[9].split(sep=' ')

# Funcion para seleccionar todas las palabras con mas de 5 letras
def getMayoresCinco(parrafo):
    palabrasM = []
    for x in parrafo:
        if len(x) >= 5:
            palabrasM.append(x)
    return palabrasM

# Funcion para eliminar Repetidos de una lista
def getUnicos(lista):
    nueva_lista = []
    for elemento in lista:
        if not elemento in nueva_lista:
            nueva_lista.append(elemento)
    return nueva_lista

print(p)
l = getMayoresCinco(p)    
q = getUnicos(l)
print(l)
print(q)
print(len(l))
print(len(q))


# Funcion para selecionar las palabras
def getSeleccion(listado):
    seleccion = []
    valor = sorted(sample(range(0 , len(listado)-1),3))
    for i in range(3):
        seleccion.append(listado[valor[i]])
    return seleccion
o = getSeleccion(q)        
print(o)


filename = "1.docx"
doc = Document(filename)
# # Funcion para modificar docx con valores seleccionados
def getMarca(lista_palabras):
    for i in range(3):
        regex1 = re.compile(lista_palabras[i])
        nuevo = lista_palabras[i] + " "
        docx_replace_regex(doc, regex1 , nuevo)
    doc.save('7.docx')
    return None

getMarca(o)

""" print(regex1)
v = " "
j = o[1] + v
replace1 = j



docx_replace_regex(doc, regex1 , replace1) """



