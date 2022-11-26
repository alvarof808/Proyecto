""" import docx

doc = docx.Document('1.docx')

x = doc.paragraphs[0].text
y = str(x)

y.replace("E", "Hola",10)
print(y)
doc.save("6.docx") """
# *******

import re
from docx import Document

def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)
filename = "1.docx"

regex1 = re.compile(r"caballería")
replace1 = r"Python34, "

#regex2 = re.compile(r"KBQ648E")
#replace2 = r"KLM789J"

doc = Document(filename)

docx_replace_regex(doc, regex1 , replace1)
#docx_replace_regex(doc, regex2 , replace2)


doc.save('7.docx')