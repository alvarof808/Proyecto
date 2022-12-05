

from cgi import print_arguments
from random import sample
# Libreria Regex para insertar valores en docx
import re
# Libreria Python-docx para manipular archivos .docx
import docx
#from docx import Document
# Libreria para convertir Pdf
import subprocess
import PyPDF2






# Funcion para modificar el archivos .docx
def docx_replace_regex(doc_obj, regex , replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text
                    
# Funcion de prueba                    
def docx_replace_regex2(doc_obj, regex , replace):
    #for p in doc_obj.paragraphs:
    #p = doc_obj.paragraphs[9]
    p = doc_obj
    if regex.search(p.text):
        inline = p.runs
        # Loop added to work with runs (strings with same style)
        for i in range(len(inline)):
            ##if regex.search(inline[i].text):
                text = regex.sub(replace, inline[i].text)
                inline[i].text = text
    

# Funcion para obtener todo el texto del documento .docx
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    x = '\n'.join(fullText)
    return x
    #return '\n'.join(fullText)

# Nro 1
# Funcion para obtener los parrafos independientes 
def getParrafo(filename):
    doc = docx.Document(filename)
    parrafo = []
    x = len(doc.paragraphs)
    for y in range(0, x):
        parrafo.append(doc.paragraphs[y].text)
    return parrafo  
  

# Nro 2
# Funcion para separar las palabras del parrafo
def getPalabras(listado):
    lista = []
    for parrafo in listado:
        lista.append(parrafo.split(sep=' '))
    #print(lista)
    return lista


# Nro 3
# Funcion para seleccionar todas las palabras con mas de 5 letras
def getMayoresCinco(lista):
    listaM = []
    for parrafo in lista:
        palabrasM = []
        for x in parrafo:
            if len(x) >= 5:
                palabrasM.append(x)
        listaM.append(palabrasM)
    return listaM

# Nro 4
# Funcion para eliminar Repetidos de una lista
def getUnicos(listado):
    listas = []
    for lista in listado:
        nueva_lista = []
        for elemento in lista:
# Este if verifica que el listado no tenga los siguientes carácteres ( ó ) ó ""
            if elemento.find(")") == -1 and elemento.find("(") == -1 and elemento.find('"')== -1 and elemento.find("á") == -1 and elemento.find("é") == -1 and elemento.find("í") == -1 and elemento.find("ó") == -1 and elemento.find("ú") == -1 and elemento.find("ñ") == -1:
            # Verificamos que no existan repetidos en lista
                #if not elemento in nueva_lista:
                nueva_lista.append(elemento)
        
        nueva_listas = getEliminarRepetidos(nueva_lista)            
        listas.append(nueva_listas)
    return listas

# Funcion para encontrar repetidos y eliminarlos
def getEliminarRepetidos(lista):
    nueva = lista
    dobles = set()
    #print(len(lista))
    #print(lista)
    #print("***********************")
    duplicados = list({x for x in nueva if x in dobles or (dobles.add(x) or False)})
   
    nueva = list(dict.fromkeys(lista))
    #nueva = list(nueva)
    for i in duplicados:
            nueva.remove(i)
    #print(len(nueva))
    #print(nueva)
    return nueva
    

# Nro 5
# Funcion para selecionar las palabras
def getSeleccion(listas):
    lista_seleccion = []
    for listado in listas:
        seleccion = []
        if len(listado) >= 6:
            valor = sorted(sample(range(0 , len(listado)-1),3))
            #print(valor)
            for i in range(3):
                seleccion.append(listado[valor[i]])
        lista_seleccion.append(seleccion)
    c = int(len(lista_seleccion))
    lista_selecciones = []
    for x in range(c):
        if len(lista_seleccion[x]) == 3:
            lista_selecciones.append(lista_seleccion[x])
    
    print(lista_seleccion)
    return lista_seleccion

# Nro 6
# Funcion para modificar docx con valores seleccionados
def getMarca(lista_palabras, doc):
    for z in range(len(doc.paragraphs)):
        if len(lista_palabras[z]) == 3:
            for i in range(3):
                regex1 = re.compile(lista_palabras[z][i])
                nuevo = lista_palabras[z][i] + ' '
                docx_replace_regex2(doc.paragraphs[z], regex1 , nuevo)
    getListTxt(lista_palabras)
    doc.save('10.docx')
    getPDF('10.docx')
    return None

# Funcion para convertir .docx a .pdf
def getPDF(filename):
    subprocess.run(["lowriter", "--convert-to", "pdf", filename])
    #pdfFileObj = open(filename[:-5]+'.pdf', 'rb')
    #pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
    #pageObj = pdfReader.getPage(0)
    #print(pageObj.extractText())
    #print()
    #subprocess.run(["rm", filename[:-5] +".pdf"])
    return None
# Funcion guardar en un txt
def getListTxt(lista):
    txt = open("test1.txt", "w")
    txt.write(str(lista))
    txt.close()
    return txt


doc = docx.Document('1.docx')


#print(f=getText('1.docx'))
#print(f[0:30])
#print(dir(doc))
#getSeccion('1.docx')
#print(x)
w  = getParrafo('1.docx') 
p = getPalabras(w)

#print(p)
l = getMayoresCinco(p)
u = getUnicos(l)     
o = getSeleccion(u)  


#print(getMayoresCinco(p))
#print(getSeleccion(l))

#l = getMayoresCinco(p)    
#filename = '1.docx'
#doc = Document(filename)
getMarca(o,doc)
