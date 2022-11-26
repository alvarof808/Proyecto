from docx import Document

#document = Document()

#f = open('1.docx', 'rb')
#document = Document(f)
#print(document)
#f.close()

# or

document = Document('1.docx')
document.save('2.docx')